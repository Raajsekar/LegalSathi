# backend/app.py
from flask import Flask, request, jsonify, send_file, Response
import re
from docx import Document
from io import BytesIO
import uuid
from flask_cors import CORS
from flask import request
from dotenv import load_dotenv
import os, uuid, time, urllib.parse, traceback
from bson.objectid import ObjectId
from flask import stream_with_context
import json
import time
from legal_engine import (
    detect_legal_intent, detect_jurisdiction, detect_writing_style,
    generate_contract, generate_tax_reply, generate_docx_stream,
    clause_review, precedent_search, make_jurisdiction_note
)

from db import db
DOCX_TEMP = {}
# AI SDK import (Groq)
try:
    from groq import Groq
except Exception:
    Groq = None

# file utils and extractors
from pdf_utils import text_to_pdf
import fitz  # pymupdf
import docx
from pymongo import MongoClient

from bson.errors import InvalidId

def is_valid_objectid(value: str):
    try:
        ObjectId(value)
        return True
    except Exception:
        return False

# Load environment
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
MONGODB_URI = os.getenv("MONGODB_URI", "")
FRONTEND_ORIGIN = os.getenv("FRONTEND_ORIGIN", "*")

# App init
app = Flask(__name__)
CORS(app, origins=[FRONTEND_ORIGIN])

# Validate Groq
if Groq is None:
    print("Warning: groq SDK not installed or import failed. Install and configure the groq package.")
else:
    try:
        client = Groq(api_key=GROQ_API_KEY)
    except Exception as e:
        print("Warning: could not initialize Groq client:", e)
        client = None

# Validate / connect Mongo (defensive)
mongo = None
db = None
chats = None
if MONGODB_URI:
    try:
        mongo = MongoClient(MONGODB_URI, serverSelectionTimeoutMS=5000)
        mongo.admin.command("ping")
        db = mongo.get_database("legalsathi")
        chats = db.get_collection("chats")
    except Exception as e:
        print("MongoDB connection warning:", e)
        mongo = db = chats = None
else:
    print("Warning: MONGODB_URI not provided. History will not be saved.")

# Prevent boolean truth testing error
if db is not None and chats is None:
    chats = db.get_collection("chats")


UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs("generated_pdfs", exist_ok=True)

# Basic GST calculation helper:
def calculate_gst(amount: float, rate_percent: float, inclusive=False, interstate=False):
    """
    Returns dict with base_amount, gst_amount, cgst, sgst, igst, total.
    - If inclusive=True, amount is GST-inclusive, we compute base and tax.
    - interstate=True -> IGST applied, else split into CGST/SGST halves.
    """
    r = float(rate_percent or 0.0)
    if inclusive:
        base = amount / (1 + r/100)
        gst_amount = amount - base
    else:
        base = amount
        gst_amount = base * r / 100.0

    if interstate:
        igst = gst_amount
        cgst = 0.0
        sgst = 0.0
    else:
        igst = 0.0
        cgst = gst_amount / 2.0
        sgst = gst_amount / 2.0

    total = base + gst_amount
    return {
        "base_amount": round(base, 2),
        "gst_amount": round(gst_amount, 2),
        "cgst": round(cgst, 2),
        "sgst": round(sgst, 2),
        "igst": round(igst, 2),
        "total_amount": round(total, 2),
        "rate_percent": r,
        "inclusive": inclusive,
        "interstate": interstate,
    }


def trim_messages(messages, max_chars=8000):
    """
    Ensure the message history does not exceed max_chars.
    We trim oldest messages first.
    """
    total = sum(len(m["content"]) for m in messages)

    if total <= max_chars:
        return messages

    trimmed = []
    running = 0

    # keep only the most recent messages
    for m in reversed(messages):
        length = len(m["content"])
        if running + length <= max_chars:
            trimmed.append(m)
            running += length
        else:
            break

    # restore chronological order
    return list(reversed(trimmed))


# -------------------------------
# CHATGPT-STYLE CONVERSATION BUILDER
# -------------------------------
def ls_build_context(conv_id, limit=30):
    """
    Loads messages in correct chronological order.
    Includes document system messages.
    """
    msgs = list(
        db.get_collection("messages")
        .find({"conv_id": ObjectId(conv_id)})
        .sort("timestamp", 1)
    )

    formatted = []
    for m in msgs:
        formatted.append({
            "role": m["role"],
            "content": m["content"]
        })

    # keep only last X messages (ChatGPT behaviour)
    if len(formatted) > limit:
        formatted = formatted[-limit:]

    return formatted

# --- AI helper ---
def ask_ai(context, prompt):
    """
    Safe wrapper around Groq. Ensures we always send at least 1 message.
    Trims content defensively and returns friendly errors if Groq/client missing.
    """
    try:
        if not str(prompt or "").strip():
            return "⚠️ No input"

        if client is None:
            print("AI client not configured.")
            return "⚠️ AI not configured"

        sys = (context or "You are LegalSathi.").strip()

        # build messages list
        messages = []
        # always include a system message
        messages.append({"role": "system", "content": sys})

        # user content must be present
        user_content = str(prompt).strip()
        messages.append({"role": "user", "content": user_content})

        # defensive trim (character-based) to avoid very large payloads
        total_len = sum(len(m["content"]) for m in messages)
        if total_len > 8000:
            # trim user content if needed
            allowed = 8000 - len(sys)
            if allowed <= 0:
                messages[1]["content"] = user_content[:2000]
            else:
                messages[1]["content"] = user_content[:allowed]

        # final defensive check: ensure messages has at least one element for Groq
        if not messages:
            messages = [{"role": "system", "content": "You are LegalSathi."},
                        {"role": "user", "content": "No input"}]

        completion = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=messages
        )
        return completion.choices[0].message.content.strip()
    except Exception as e:
        print("Groq Error:", e)
        traceback.print_exc()
        return "⚠️ Sorry, the AI faced an error. Please try again."

# --- File Text Extractors ---
def extract_pdf_text(filepath):
    text = ""
    with fitz.open(filepath) as pdf:
        for page in pdf:
            text += page.get_text("text") + "\n"
    return text.strip()


def extract_docx_text(filepath):
    doc = docx.Document(filepath)
    return "\n".join([p.text for p in doc.paragraphs])

# --- Conversation helpers ---
def create_conversation(user_id, title="New conversation"):
    conv = {
        "user_id": user_id,
        "title": title,
        "created_at": time.time(),
        "updated_at": time.time()
    }
    res = db.get_collection("conversations").insert_one(conv)
    conv["_id"] = str(res.inserted_id)
    return conv

def add_message(conv_id, role, content):
    msg = {
        "conv_id": ObjectId(conv_id),
        "role": role,
        "content": content,
        "timestamp": time.time()
    }
    db.get_collection("messages").insert_one(msg)

def build_context(conv_id, max_messages=12):
    """
    Returns list of dicts for system/user messages suitable to include with AI call.
    We'll include last N messages from the conversation.
    """
    msgs = list(db.get_collection("messages")
                .find({"conv_id": ObjectId(conv_id)})
                .sort("timestamp", -1)
                .limit(max_messages))
    msgs = list(reversed(msgs))
    # convert to simple list:
    return [{"role": m["role"], "content": m["content"]} for m in msgs]

def simulate_stream(text, chunk_size=30, delay=0.03):
    """
    Yield the text in small chunks (fallback if AI SDK has no streaming).
    chunk_size is characters per chunk.
    """
    i = 0
    while i < len(text):
        chunk = text[i:i+chunk_size]
        i += chunk_size
        time.sleep(delay)  # short pause so frontend gets streaming feel
        yield chunk

# ============================================================
#  GLOBAL LEGAL INTENT ROUTER + JURISDICTION DETECTOR + STYLE ENGINE
# ============================================================




# -------------------------------------------
# 1. AUTO JURISDICTION DETECTOR (GLOBAL)
# -------------------------------------------
def detect_jurisdiction(text: str):
    t = text.lower()

    # USA STATES
    usa_states = {
        "california", "texas", "new york", "florida", "illinois",
        "ohio", "georgia", "pennsylvania", "washington", "virginia"
    }

    for st in usa_states:
        if st in t:
            return ("USA", st.title())

    # COUNTRY DETECTION
    if "india" in t or "gst" in t or "gstr" in t:
        return ("India", "India")

    if "dubai" in t or "uae" in t or "sharjah" in t or "fta" in t:
        return ("UAE", "UAE")

    if "singapore" in t:
        return ("Singapore", "Singapore")

    if "uk" in t or "england" in t or "wales" in t or "hmrc" in t:
        return ("UK", "United Kingdom")

    if "australia" in t or "nsw" in t or "queensland" in t:
        return ("Australia", "Australia")

    if "canada" in t:
        return ("Canada", "Canada")

    if "europe" in t or "eu" in t or "gdpr" in t:
        return ("EU", "European Union")

    # Default global
    return ("Global", "Generic")


# -------------------------------------------
# 2. WRITING STYLE DETECTOR (Simple vs Legal English)
# -------------------------------------------
def detect_writing_style(message: str):
    """
    Auto-detect style:
    - Simple English for common users
    - Legal English for professional/legal queries
    """

    msg = message.lower()

    legal_keywords = [
        "whereas", "hereto", "hereby", "indemnity",
        "governing law", "jurisdiction", "arbitration",
        "non-disclosure", "breach", "termination clause",
        "section", "sub-section", "pursuant", "notwithstanding",
        "liability", "force majeure"
    ]

    if any(word in msg for word in legal_keywords):
        return "legal"

    if len(message.split()) < 6:
        return "simple"

    # Tax & notice replies use simple format by default
    tax_flags = ["gst", "vat", "irs", "hmrc", "fta", "143", "notice", "scn", "drc"]
    if any(w in msg for w in tax_flags):
        return "simple"

    # Default
    return "simple"


# -------------------------------------------
# 3. GLOBAL LEGAL INTENT ROUTER
# -------------------------------------------
def detect_legal_intent(message: str):
    """
    Returns one of:
    - 'contract'
    - 'notice_reply'
    - 'tax_reply'
    - 'clause_review'
    - 'document_summary'
    - 'lawyer_mode'
    - 'generic_chat'
    """

    m = message.lower()

    # CONTRACT / AGREEMENT DRAFTING
    contract_terms = [
        "draft agreement", "draft contract", "prepare agreement", "create contract",
        "employment agreement", "rental agreement", "partnership deed",
        "mou", "nda", "service agreement", "lease agreement"
    ]
    if any(t in m for t in contract_terms):
        return "contract"

    # TAX REPLY
    tax_patterns = [
        r"\bscn\b", r"drc-01", r"asmt-10", r"143\(1\)", r"143\(2\)", r"\b148\b",
        "gst notice", "vat notice", "hmrc", "irs", "cp2000", "fta notice"
    ]
    if any(re.search(p, m) for p in tax_patterns):
        return "tax_reply"

    # GENERAL NOTICE / LEGAL REPLY
    if "legal notice" in m or "reply notice" in m or "send a notice" in m:
        return "notice_reply"

    # CLAUSE REVIEW
    review_terms = ["review clause", "improve clause", "rewrite clause", "redraft clause"]
    if any(t in m for t in review_terms):
        return "clause_review"

    # DOCUMENT SUMMARY
    if "summarize" in m or "highlight points" in m or "explain this document" in m:
        return "document_summary"

    # LAWYER MODE
    if "section" in m or "supreme court" in m or "precedent" in m:
        return "lawyer_mode"

    # Default
    return "generic_chat"

# conversation_builder.py

def build_conversation(conv_id, new_user_message):
    """
    Builds a ChatGPT-style conversation history.
    System messages (document contents) are included.
    """
    messages_col = db.get_collection("messages")

    # fetch all messages in this conversation
    msgs = list(messages_col.find({"conv_id": ObjectId(conv_id)}).sort("timestamp", 1))

    formatted = []

    for m in msgs:
        role = m["role"]  # system, user, assistant
        content = m["content"]

        # pass as-is
        formatted.append({"role": role, "content": content})

    # finally append the NEW user message
    formatted.append({"role": "user", "content": new_user_message})

    return formatted

def generate_contract(jurisdiction, message, style="simple"):
    """
    Generate a global contract draft template based on jurisdiction.
    """
    country, region = jurisdiction

    # Writing style
    if style == "legal":
        tone = "Use formal legal English, contract-style sentences, numbered clauses, and professional structure."
    else:
        tone = "Use clear, simple English with easy-to-understand clauses."

    # Jurisdiction-specific governing law
    governing = {
        "India": "This Agreement shall be governed by the laws of India.",
        "UAE": "This Agreement shall be governed by the laws of the United Arab Emirates.",
        "USA": f"This Agreement shall be governed by the laws of the State of {region}, USA.",
        "UK": "This Agreement shall be governed by the laws of England and Wales.",
        "Singapore": "This Agreement shall be governed by the laws of Singapore.",
        "Australia": "This Agreement shall be governed by the laws of Australia.",
        "Canada": "This Agreement shall be governed by the laws of Canada.",
        "EU": "This Agreement shall comply with relevant EU Contract Laws and Regulations.",
        "Global": "This Agreement shall be interpreted under internationally accepted legal principles."
    }.get(country, "This Agreement shall be interpreted under internationally accepted legal principles.")

    draft = f"""
{tone}

### AGREEMENT / CONTRACT – AUTO-DRAFTED

**Jurisdiction:** {country} ({region})

---

### 1. Parties
This Agreement is made between:
- **Party A:** ______________________  
- **Party B:** ______________________  

---

### 2. Purpose
Describe the purpose of this contract clearly:
“{message[:300]}...”

---

### 3. Term
The Agreement shall commence on __________ and remain valid until __________ unless terminated earlier.

---

### 4. Responsibilities of Party A
- ____________________________
- ____________________________
- ____________________________

### 5. Responsibilities of Party B
- ____________________________
- ____________________________
- ____________________________

---

### 6. Payment Terms
- Amount payable: __________  
- Currency: __________  
- Payment due on: __________  
- Method of payment: __________  

---

### 7. Confidentiality
Both parties agree to maintain strict confidentiality regarding all shared information.

---

### 8. Intellectual Property
All intellectual property created under this Agreement shall belong to:
□ Party A  
□ Party B  
□ Jointly (tick one)

---

### 9. Indemnity
Each party agrees to indemnify the other against losses arising from breach, negligence, or misconduct.

---

### 10. Limitation of Liability
Neither party shall be liable for indirect, incidental, or consequential damages.

---

### 11. Termination
This Agreement may be terminated:
- By mutual agreement  
- For breach of terms  
- Upon written notice of ____ days  

---

### 12. Dispute Resolution
Disputes shall be resolved through:
□ Negotiation  
□ Mediation  
□ Arbitration  
□ Court proceedings

---

### 13. Governing Law
{governing}

---

### 14. Signatures

Party A: _____________________  Date: ________

Party B: _____________________  Date: ________

---

*This is an auto-generated draft. Please review before use.*
"""

    return draft.strip()


# -------------------------------------------
# 5. TAX NOTICE REPLY ENGINE (GLOBAL)
# -------------------------------------------
def generate_tax_reply(jurisdiction, message, style="simple"):
    """
    Automatically generate a tax notice reply for GST/VAT/IRS/HMRC/Income Tax.
    """
    country, region = jurisdiction

    intro = {
        "India": "Subject: Reply to GST / Income Tax Notice",
        "UAE": "Subject: Response to UAE FTA VAT Notice",
        "UK": "Subject: Response to HMRC VAT Compliance Notice",
        "USA": "Subject: Response to IRS Notice (including CP2000)",
        "EU": "Subject: Response to EU Tax Compliance Communication",
        "Global": "Subject: Response to Tax / Compliance Notice"
    }.get(country, "Subject: Reply to Tax Notice")

    if style == "legal":
        tone = "Use formal language suitable for tax authorities, referencing relevant statutes where appropriate."
    else:
        tone = "Explain clearly in simple English without legal jargon."

    reply = f"""
{intro}

{tone}

---

### 1. Basic Reference Details
- Notice Reference Number: _____________________
- Date of Notice: _____________________
- Tax Period: _____________________

---

### 2. Acknowledgement
We acknowledge receipt of your notice regarding:

“{message[:250]}...”

---

### 3. Background / Facts
Provide brief facts:
- Nature of business  
- Relevant transactions  
- Key documents already submitted  

---

### 4. Clarification / Explanation
Insert specific explanation for each point raised in the notice:
- Issue 1: _____________________  
- Issue 2: _____________________  
- Issue 3: _____________________  

---

### 5. Supporting Evidence
We are enclosing the following:
- Invoices  
- Bank statements  
- Ledgers  
- Contracts  
- Other supporting documents  

---

### 6. Legal/Procedural Position
(If applicable based on jurisdiction)
- GST Act (India)  
- VAT Decree (UAE)  
- HMRC VAT Rules (UK)  
- IRS Code (USA)  

---

### 7. Conclusion
We request the authority to kindly consider the above explanation and drop the proceedings.

---

### 8. Declaration
We declare that the information furnished above is true and correct.

---

Authorized Signatory  
Name: _____________________  
Date: _____________________

---

*This is an auto-generated reply. Cross-check before submission.*
"""

    return reply.strip()


# -------------------------------------------
# 6. PDF + DOCX GENERATION
# -------------------------------------------
def generate_docx_stream(text: str):
    """
    Create a DOCX file in memory and return (filename, BytesIO buffer)
    """
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"{uuid.uuid4().hex}.docx"
    return filename, buffer
# --- Routes ---

@app.route("/api/gst/calc", methods=["POST"])
def api_gst_calc():
    """
    POST JSON:
      { "amount": 1000, "rate": 18, "inclusive": false, "interstate": false }
    Returns GST calculation details.
    """
    data = request.get_json(force=True)
    try:
        amount = float(data.get("amount", 0))
        rate = float(data.get("rate", 18))
        inclusive = bool(data.get("inclusive", False))
        interstate = bool(data.get("interstate", False))
    except Exception:
        return jsonify({"error": "Invalid input"}), 400

    result = calculate_gst(amount, rate, inclusive=inclusive, interstate=interstate)
    return jsonify(result)

@app.route("/api/gst/tips")
def api_gst_tips():
    """
    Return a short, safe list of lawful GST/tax planning tips and resources.
    These are general pointers — not professional advice.
    """
    tips = [
        {
            "title": "Choose correct HSN/SAC & Invoice format",
            "description": "Use correct HSN/SAC codes and maintain tax invoices — critical for input tax credit claims."
        },
        {
            "title": "Claim Input Tax Credit (ITC) properly",
            "description": "Maintain GST-compliant invoices and reconcile GSTR2B/2A to avoid blocked credits. Check composition scheme thresholds before opting in."
        },
        {
            "title": "Composition scheme vs regular registration",
            "description": "Small businesses may opt for composition (lower compliance) but composition dealers cannot claim ITC — choose based on business model."
        },
        {
            "title": "Keep records for 6 years",
            "description": "Maintain invoices, E-way bills and GST returns for statutory retention (subject to updates in law)."
        },
        {
            "title": "When in doubt, consult a CA",
            "description": "GST law is complex; for planning/loopholes consult a qualified Chartered Accountant — our tips are educational only."
        },
    ]
    return jsonify(tips)
@app.route("/")
def home():
    return "⚖️ LegalSathi backend active"

@app.route("/api/clause_review", methods=["POST"])
def api_clause_review():
    data = request.get_json(force=True)
    clause = data.get("clause", "")
    if not clause:
        return jsonify({"error": "Missing clause"}), 400
    res = clause_review(clause)
    return jsonify(res)

@app.route("/api/precedent_search", methods=["GET"])
def api_precedent_search():
    user_id = request.args.get("user_id")
    q = request.args.get("q", "")
    hits = precedent_search(user_id, q, limit=8)
    return jsonify(hits)


@app.route("/download_docx/<filename>")
def download_docx(filename):
    if filename not in DOCX_TEMP:
        return "File not found", 404

    buffer = DOCX_TEMP[filename]
    return send_file(
        buffer,
        download_name=filename,
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
# ============================================================
#  PART 3 — FULL UPGRADED /api/stream_chat (GLOBAL LEGALSATHI ENGINE)
# ============================================================

@app.route("/api/stream_chat", methods=["POST"])
def stream_chat():
    """
    GLOBAL LEGALSATHI STREAMING ENGINE
    - Detects legal intent
    - Detects jurisdiction
    - Detects writing style
    - Uses drafting & tax engines
    - Streams AI output
    - After stream completes → generates PDF + DOCX
    """

    try:
        data = request.get_json(force=True)
        user_id = data.get("user_id")
        message = (data.get("message") or "").strip()
        conv_id = data.get("conv_id")

        if not user_id or not message:
            return jsonify({"error": "Missing user_id or message"}), 400

        # -----------------------------------------------------
        # 1. CREATE OR VALIDATE CONVERSATION
        # -----------------------------------------------------
        conv_doc = None

        # New conversation?
        if not conv_id or not is_valid_objectid(conv_id):
            conv = create_conversation(user_id, title="New conversation")
            conv_id = conv["_id"]
            conv_doc = conv

        else:
            # Validate existing conversation (if DB exists)
            if db is not None:
                try:
                    conv_doc = db.get_collection("conversations").find_one({"_id": ObjectId(conv_id)})
                    if not conv_doc or conv_doc.get("user_id") != user_id:
                        return jsonify({"error": "Invalid conversation ID"}), 403
                except:
                    return jsonify({"error": "Invalid conversation ID"}), 403
            else:
                conv_doc = {"_id": conv_id, "user_id": user_id}

        # Save user message
        add_message(conv_id, "user", message)

        # -----------------------------------------------------
        # 2. DETECT INTENT, JURISDICTION, STYLE
        # -----------------------------------------------------
        intent = detect_legal_intent(message)
        jurisdiction = detect_jurisdiction(message)
        style = detect_writing_style(message)

        # -----------------------------------------------------
        # 3. SYSTEM PROMPT BASED ON ROUTED INTENT
        # -----------------------------------------------------
        if intent == "contract":
            system_prompt = "You are LegalSathi, a global contract drafting expert. Generate structured legal agreements as per detected jurisdiction."

        elif intent == "tax_reply":
            system_prompt = "You are LegalSathi, a global tax notice reply expert. Generate clear, concise, structured replies for GST, VAT, IRS, HMRC, IT notices."

        elif intent == "notice_reply":
            system_prompt = "You are LegalSathi, draft professional legal notices and replies."

        elif intent == "clause_review":
            system_prompt = "You are LegalSathi, a clause rewriting expert. Improve, polish, and legally strengthen clauses."

        elif intent == "document_summary":
            system_prompt = "You are LegalSathi, summarize documents clearly with risks highlighted."

        elif intent == "lawyer_mode":
            system_prompt = "You are LegalSathi-ADV, behaving like a senior lawyer. Provide deep legal reasoning, citations, and structured guidance."

        else:
            system_prompt = "You are LegalSathi, a global AI legal assistant."

        # Branding only ONCE (first assistant message)
        message_history = ls_build_context(conv_id, limit=30)

        is_first_reply = len(message_history) <= 1

        branding = ""
        if is_first_reply:
            branding = (
                "I am **LegalSathi**, your global AI legal assistant. "
                "I can draft agreements, notices, tax replies, and legal documents for any country.\n\n"
            )

        # -----------------------------------------------------
        # 4. PRE-GENERATE TEMPLATE IF CONTRACT or TAX ENGINE
        # -----------------------------------------------------
        pre_generated = None

        if intent == "contract":
            pre_generated = generate_contract(jurisdiction, message, style)

        elif intent == "tax_reply":
            pre_generated = generate_tax_reply(jurisdiction, message, style)

        # Add branding + pre-generated structure to messages
        user_message_payload = branding
        if pre_generated:
            user_message_payload += f"Here is the structured draft template:\n\n{pre_generated}\n\nNow refine the above draft based on the user's request:\n\n{message}"
        else:
            user_message_payload += message

        # -----------------------------------------------------
        # 5. BUILD AI MESSAGE LIST
        # -----------------------------------------------------
        messages_for_ai = [{"role": "system", "content": system_prompt}]
        messages_for_ai.extend(message_history)
        messages_for_ai.append({"role": "user", "content": user_message_payload})

        messages_for_ai = trim_messages(messages_for_ai, max_chars=7500)

        # -----------------------------------------------------
        # 6. STREAMING RESPONSE GENERATOR
        # -----------------------------------------------------
        def generate_stream():
            final_text = ""

            # Stream from GROQ
            try:
                for chunk in client.chat.completions.create(
                    model="llama-3.1-8b-instant",
                    messages=messages_for_ai,
                    stream=True
                ):
                    if (
                        chunk
                        and chunk.choices
                        and hasattr(chunk.choices[0].delta, "content")
                        and chunk.choices[0].delta.content
                    ):
                        delta = chunk.choices[0].delta.content
                        final_text += delta
                        yield json.dumps({"chunk": delta}) + "\n"

            except Exception as stream_err:
                print("Streaming failed:", stream_err)
                yield json.dumps({"chunk": "\n[Streaming failed]\n"}) + "\n"

            # -----------------------------------------------------
            # 7. SAVE ASSISTANT RESPONSE
            # -----------------------------------------------------
            try:
                # ChatGPT-style safe write: no overwrite, no merge issues
                db.get_collection("messages").insert_one({
    "conv_id": ObjectId(conv_id),
    "role": "assistant",
    "content": final_text,
    "timestamp": time.time()
})

                db.get_collection("conversations").update_one(
                    {"_id": ObjectId(conv_id)},
                    {
                        "$set": {
                            "updated_at": time.time(),
                            "title": (final_text[:60] or "Conversation"),
                            "last_message": final_text
                        }
                    }
                )
            except Exception as e:
                print("Save assistant error:", e)

            # -----------------------------------------------------
            # 8. GENERATE PDF + DOCX AFTER STREAM ENDS
            # -----------------------------------------------------
            from pdf_utils import text_to_pdf
            pdf_name = f"{uuid.uuid4().hex}.pdf"
            pdf_path = text_to_pdf(final_text, pdf_name)

            # DOCX (in-memory)
            docx_name, docx_buffer = generate_docx_stream(final_text)

            # Save DOCX buffer to temp memory store
            DOCX_TEMP[docx_name] = docx_buffer

            # -----------------------------------------------------
            # 9. FINAL SIGNAL
            # -----------------------------------------------------
            yield json.dumps({
                "done": True,
                "conv_id": conv_id,
                "pdf_url": f"/download/{pdf_name}",
                "docx_url": f"/download_docx/{docx_name}"
            }) + "\n"

        return Response(stream_with_context(generate_stream()), mimetype="text/plain")

    except Exception as e:
        print("stream_chat error:", e)
        traceback.print_exc()
        return jsonify({"error": "internal server error"}), 500

@app.route("/api/conversations/<user_id>")
def get_conversations(user_id):
    try:
        convs = list(db.get_collection("conversations").find({"user_id": user_id}).sort("updated_at", -1))
        for c in convs:
            c["_id"] = str(c["_id"])
        return jsonify(convs)
    except Exception as e:
        print("get_conversations error:", e)
        return jsonify([])

@app.route("/api/conversation/<conv_id>")
def get_conversation(conv_id):
    try:
        msgs = list(db.get_collection("messages").find({"conv_id": ObjectId(conv_id)}).sort("timestamp", 1))
        out = []
        for m in msgs:
            m["_id"] = str(m["_id"])
            m["conv_id"] = str(m["conv_id"])
            out.append(m)
        return jsonify(out)
    except Exception as e:
        print("get_conversation error:", e)
        return jsonify([])

@app.route("/api/files/<user_id>")
def list_files(user_id):
    try:
        docs = list(db.get_collection("file_records").find({"user_id": user_id}).sort("timestamp", -1))
        for d in docs:
            d["_id"] = str(d["_id"])
        return jsonify(docs)
    except Exception as e:
        print("list_files error:", e)
        return jsonify([])
@app.route("/api/upload", methods=["POST"])
def upload_file():
    """
    ChatGPT-style upload:
    - Does NOT create a new conversation unless conv_id is missing
    - Stores extracted document text as a hidden assistant message
    - Allows MULTI-TURN chat with the file
    - Lets users ask followup questions about the PDF
    """
    try:
        user_id = request.form.get("user_id")
        conv_id = request.form.get("conv_id")  # new: continue same chat
        task = request.form.get("task", "summarize")
        file = request.files.get("file")

        if not user_id or not file:
            return jsonify({"error": "Missing fields"}), 400

        # ---- Save actual uploaded file ----
        filename = f"{uuid.uuid4().hex}_{file.filename}"
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)

        # ---- Extract text from uploaded doc ----
        lower = filename.lower()
        if lower.endswith(".pdf"):
            content = extract_pdf_text(filepath)
        elif lower.endswith(".docx"):
            content = extract_docx_text(filepath)
        elif lower.endswith(".txt"):
            content = open(filepath, "r", encoding="utf8", errors="ignore").read()
        else:
            return jsonify({"error": "Unsupported file type"}), 400

        # ---- Create conversation if needed ----
        if not conv_id or not is_valid_objectid(conv_id):
            conv = {
                "user_id": user_id,
                "title": file.filename,
                "created_at": time.time(),
                "updated_at": time.time(),
                "last_message": ""
            }
            conv_res = db.get_collection("conversations").insert_one(conv)
            conv_id = str(conv_res.inserted_id)

        # ---- (KEY FEATURE) Insert the file content as a hidden system message ----
        # This is how ChatGPT allows follow-up questions about the file.
        db.get_collection("messages").insert_one({
            "conv_id": ObjectId(conv_id),
            "role": "system",
            "content": f"__FILE_CONTENT__\n{content}",
            "timestamp": time.time()
        })

        # ---- Save user message indicating upload ----
        db.get_collection("messages").insert_one({
            "conv_id": ObjectId(conv_id),
            "role": "user",
            "content": f"📄 Uploaded file: {file.filename}",
            "timestamp": time.time()
        })

        # ---- Generate AI summary (first response) ----
        summary_prompt = (
            "You are LegalSathi. The user uploaded a document. "
            "Provide a clear high-level summary. "
            "DO NOT lose the document content — future prompts will reference it."
        )
        reply = ask_ai(summary_prompt, content[:8000])

        # ---- Save assistant reply ----
        db.get_collection("messages").insert_one({
            "conv_id": ObjectId(conv_id),
            "role": "assistant",
            "content": reply,
            "timestamp": time.time()
        })

        # ---- Update conversation metadata ----
        db.get_collection("conversations").update_one(
            {"_id": ObjectId(conv_id)},
            {"$set": {"updated_at": time.time(), "last_message": reply}}
        )

        # ---- Generate summary PDF ----
        pdfname = f"{uuid.uuid4().hex[:8]}.pdf"
        text_to_pdf(reply, pdfname)

        # ---- Also save file record ----
        db.get_collection("file_records").insert_one({
            "user_id": user_id,
            "original_name": file.filename,
            "stored_path": filepath,
            "pdf": pdfname,
            "conv_id": ObjectId(conv_id),
            "timestamp": time.time()
        })

        return jsonify({
            "reply": reply,
            "pdf_url": f"/download/{pdfname}",
            "conv_id": str(conv_id)
        })

    except Exception as e:
        print("upload error:", e)
        return jsonify({"error": "internal error"}), 500


@app.route("/download/<filename>")
def download(filename):
    path = os.path.join("generated_pdfs", filename)
    if os.path.exists(path):
        return send_file(path, as_attachment=True)
    return "File not found", 404


@app.route("/api/library/<user_id>")
def library(user_id):
    """Return uploaded document summaries for that user"""
    try:
        if chats is None:
            return jsonify([])
        files = list(chats.find(
            {"user_id": user_id, "file_name": {"$exists": True}}
        ).sort("timestamp", -1))
        for f in files:
            f["_id"] = str(f["_id"])
        return jsonify(files)
    except Exception as e:
        print("API /api/library error:", e)
        traceback.print_exc()
        return jsonify([])


@app.route("/api/conversation/<conv_id>", methods=["DELETE"])
def delete_conversation(conv_id):
    try:
        db.get_collection("conversations").delete_one({"_id": ObjectId(conv_id)})
        db.get_collection("messages").delete_many({"conv_id": ObjectId(conv_id)})
        return jsonify({"status": "ok"})
    except Exception as e:
        print("delete_conversation error:", e)
        return jsonify({"error": "failed"}), 500

@app.route("/api/delete_conversation/<conv_id>", methods=["DELETE"])
def delete_conversation_alias(conv_id):
    return delete_conversation(conv_id)

@app.errorhandler(500)
def internal_error(error):
    print("500 error:", error)
    return jsonify({"error": "Internal server error"}), 500

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    print(f"🚀 LegalSathi backend is running on 0.0.0.0:{port}/")
    app.run(host="0.0.0.0", port=port)
